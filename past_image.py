#СКРИПТ ДЛЯ ВСТАВКИ БОЛЬШОГО КОЛИЧЕСТВА КАРТИНОК И ТЕКСТА В ФАЙЛ

import os
from docx import Document
from docx.shared import Inches
from PIL import Image

# Параметры скрипта
#text_file_path = 'C:/Users/Алёнка/Documents/test/СНИЛС/voyna-i-mir+snils.txt'
image_folder = 'C:/Users/Алёнка/Documents/test/IP/'
output_docx = 'C:/Users/Алёнка/Documents/test/IP/many_photos_500.docx'
n = 500 # количество раз вставки изображений

# Открытие файла с текстом
# with open(text_file_path, 'r', encoding='utf-8') as f:
#     text = f.read()

# Получение списка файлов изображений в папке
#image_formats = ['.bmp', '.pnm', '.png', '.jfif', '.jpeg', '.jpg', '.tiff']
image_formats = ['.png', '.jpg']
image_files = [f for f in os.listdir(image_folder) if
               os.path.isfile(os.path.join(image_folder, f)) and os.path.splitext(f)[1].lower() in image_formats]

# Создание нового документа Word
doc = Document()

# Вставка изображений n раз
for _ in range(n):
    for image_file in image_files:
        image_path = os.path.join(image_folder, image_file)

        try:
            img = Image.open(image_path)
            width, height = img.size

            # Вставка изображения в документ
            doc.add_picture(image_path, width=Inches(6), height=Inches(height * 6 / width))

            # Добавление разрыва страницы после каждого изображения
            doc.add_page_break()

        except Exception as e:
            print(f"Ошибка при обработке изображения '{image_path}': {e}")

# k = 100
# for _ in range(k):
#     doc.add_paragraph(text)

# Сохранение документа
doc.save(output_docx)

print(f"Документ '{output_docx}' успешно создан с вставленными изображениями.")

